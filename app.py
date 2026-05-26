import streamlit as st
import pandas as pd
import json
import uuid
from datetime import datetime
from pathlib import Path
from io import BytesIO
from docx import Document
from docx.shared import Pt
import gspread
from google.oauth2.service_account import Credentials

# ── 設定 ──────────────────────────────────────────────
DATA_FILE        = Path("data/submissions.json")
TEMPLATE_FILE    = Path("rirekisho_template.docx")
SHOKUMU_TEMPLATE = Path("shokumu_template.docx")
MAX_COMPANIES    = 5
SPREADSHEET_ID   = "1hhVNfpIHcNW-rjxcxVsD14QAPair4IIa5FYK6U9Ny_k"

# クライアントごとのパスワード設定
# st.secretsに設定がなければデモ用パスワードを使用
try:
    CLIENT_PASSWORDS = dict(st.secrets["client_passwords"])
except Exception:
    CLIENT_PASSWORDS = {
        "admin": "admin1234",
    }

def get_admin_pass():
    return "admin1234"

# ── Google Sheets接続 ──────────────────────────────────
def get_gsheet():
    try:
        creds_dict = dict(st.secrets["gcp_service_account"])
        scopes = [
            "https://www.googleapis.com/auth/spreadsheets",
            "https://www.googleapis.com/auth/drive",
        ]
        creds = Credentials.from_service_account_info(creds_dict, scopes=scopes)
        client = gspread.authorize(creds)
        sheet = client.open_by_key(SPREADSHEET_ID).sheet1
        return sheet
    except Exception as e:
        return None

def save_to_gsheet(record):
    try:
        sheet = get_gsheet()
        if sheet is None:
            return False
        # ヘッダーがなければ追加
        if sheet.row_count == 0 or sheet.cell(1, 1).value is None:
            headers = ["id","submitted_at","name","furigana_name","gender","birthday","age",
                      "postal","address","phone","email","nearest_station",
                      "dependents","spouse","spouse_support","summary","skills","pr"]
            sheet.append_row(headers)
        row = [
            record.get("id",""),
            record.get("submitted_at",""),
            record.get("name",""),
            record.get("furigana_name",""),
            record.get("gender",""),
            record.get("birthday",""),
            record.get("age",""),
            record.get("postal",""),
            record.get("address",""),
            record.get("phone",""),
            record.get("email",""),
            record.get("nearest_station",""),
            record.get("dependents",""),
            record.get("spouse",""),
            record.get("spouse_support",""),
            record.get("summary",""),
            record.get("skills",""),
            record.get("pr",""),
        ]
        sheet.append_row(row)
        return True
    except Exception as e:
        return False

TEMPLATE_FILE    = Path("rirekisho_template.docx")
SHOKUMU_TEMPLATE = Path("shokumu_template.docx")
MAX_COMPANIES    = 5

# クライアントごとのパスワード設定
# st.secretsに設定がなければデモ用パスワードを使用
try:
    CLIENT_PASSWORDS = dict(st.secrets["client_passwords"])
except Exception:
    CLIENT_PASSWORDS = {
        "admin": "admin1234",
    }

def get_admin_pass():
    return "admin1234"

def load_from_gsheet():
    try:
        sheet = get_gsheet()
        if sheet is None:
            return []
        rows = sheet.get_all_records()
        # スプレッドシートの基本情報だけなので、ローカルと統合
        return rows
    except Exception:
        return []

def save_to_gsheet_full(record):
    """全データをJSONとしてスプレッドシートに保存"""
    try:
        sheet = get_gsheet()
        if sheet is None:
            return False
        # ヘッダーがなければ追加
        try:
            header_val = sheet.cell(1, 1).value
        except Exception:
            header_val = None
        if not header_val:
            sheet.append_row(["id", "submitted_at", "name", "json_data", "client_id"])
        row = [
            record.get("id", ""),
            record.get("submitted_at", ""),
            record.get("name", ""),
            json.dumps(record, ensure_ascii=False),
            record.get("client_id", "admin"),
        ]
        sheet.append_row(row)
        return True
    except Exception as e:
        return False

def load_all_from_gsheet(client_id=None):
    """スプレッドシートから全レコードをJSONで取得（client_idでフィルタ可能）"""
    try:
        sheet = get_gsheet()
        if sheet is None:
            return []
        all_values = sheet.get_all_values()
        if len(all_values) <= 1:
            return []
        records = []
        for row in all_values[1:]:  # ヘッダーをスキップ
            if len(row) >= 4 and row[3]:
                try:
                    record = json.loads(row[3])
                    # client_idフィルタ
                    row_client = row[4] if len(row) >= 5 else "admin"
                    if client_id is None or row_client == client_id:
                        records.append(record)
                except Exception:
                    pass
        return records
    except Exception:
        return []

def delete_from_gsheet(record_id):
    """スプレッドシートから該当IDの行を削除"""
    try:
        sheet = get_gsheet()
        if sheet is None:
            return False
        all_values = sheet.get_all_values()
        for i, row in enumerate(all_values):
            if len(row) >= 1 and row[0] == record_id:
                sheet.delete_rows(i + 1)
                return True
        return False
    except Exception:
        return False

st.set_page_config(page_title="経歴入力フォーム", page_icon="📄", layout="centered")

# ── データ保存・読み込み ──────────────────────────────
def load_data():
    DATA_FILE.parent.mkdir(exist_ok=True)
    if DATA_FILE.exists():
        return json.loads(DATA_FILE.read_text(encoding="utf-8"))
    return []

def save_data(records):
    DATA_FILE.parent.mkdir(exist_ok=True)
    DATA_FILE.write_text(json.dumps(records, ensure_ascii=False, indent=2), encoding="utf-8")

# ── セルへの書き込み ──────────────────────────────────
def set_cell(cell, text):
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    if cell.paragraphs:
        run = cell.paragraphs[0].add_run(str(text) if text else "")
        run.font.size = Pt(10.5)
    else:
        cell.text = str(text) if text else ""

# ── 履歴書テンプレートに流し込む ─────────────────────
def make_rirekisho(d):
    if not TEMPLATE_FILE.exists():
        raise FileNotFoundError("rirekisho_template.docx が見つかりません。")
    buf = BytesIO()
    buf.write(TEMPLATE_FILE.read_bytes())
    buf.seek(0)
    doc = Document(buf)

    edu_list = d.get("education", [])
    car_list  = d.get("career",    [])
    lic_list  = d.get("licenses",  [])

    t0 = doc.tables[0]
    set_cell(t0.rows[0].cells[1], d.get("furigana_name",""))
    set_cell(t0.rows[1].cells[1], d.get("name",""))
    set_cell(t0.rows[1].cells[3], d.get("gender",""))
    set_cell(t0.rows[2].cells[1], f"{d.get('birthday','')}（満{d.get('age','')}歳）")
    set_cell(t0.rows[3].cells[1], d.get("furigana_address",""))
    set_cell(t0.rows[4].cells[1], f"〒{d.get('postal','')}　{d.get('address','')}")
    set_cell(t0.rows[5].cells[1], d.get("phone",""))
    set_cell(t0.rows[5].cells[5], d.get("email",""))

    t1 = doc.tables[1]
    for i, edu in enumerate(edu_list[:6]):
        r = i + 2
        if r < len(t1.rows):
            set_cell(t1.rows[r].cells[0], edu.get("year",""))
            set_cell(t1.rows[r].cells[1], edu.get("month",""))
            set_cell(t1.rows[r].cells[2], edu.get("content",""))

    for i, job in enumerate(car_list[:9]):
        r = i + 9
        if r < len(t1.rows):
            set_cell(t1.rows[r].cells[0], job.get("year",""))
            set_cell(t1.rows[r].cells[1], job.get("month",""))
            set_cell(t1.rows[r].cells[2], job.get("content",""))

    t2 = doc.tables[2]
    for i, lic in enumerate(lic_list[:3]):
        r = i + 1
        if r < len(t2.rows):
            set_cell(t2.rows[r].cells[0], lic.get("year",""))
            set_cell(t2.rows[r].cells[1], lic.get("month",""))
            set_cell(t2.rows[r].cells[2], lic.get("content",""))

    t3 = doc.tables[3]
    set_cell(t3.rows[1].cells[1], d.get("nearest_station",""))
    set_cell(t3.rows[1].cells[2], f"（配偶者を除く）{d.get('dependents','0')}人")
    set_cell(t3.rows[1].cells[3], d.get("spouse","無"))
    set_cell(t3.rows[1].cells[4], d.get("spouse_support","無"))

    t4 = doc.tables[4]
    set_cell(t4.rows[1].cells[0], d.get("pr",""))

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out

# ── 職務経歴書テンプレートに流し込む ─────────────────
def make_shokumu(d):
    if not SHOKUMU_TEMPLATE.exists():
        raise FileNotFoundError("shokumu_template.docx が見つかりません。")

    buf_in = BytesIO(SHOKUMU_TEMPLATE.read_bytes())
    doc = Document(buf_in)
    paras = doc.paragraphs

    def set_run(para, text):
        for run in para.runs:
            run.text = ""
        if para.runs:
            para.runs[0].text = str(text) if text else ""
        else:
            para.add_run(str(text) if text else "")

    def add_bold(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10.5)

    def add_normal(text):
        p = doc.add_paragraph()
        p.add_run(str(text) if text else "")

    if paras[2].runs:
        paras[2].runs[0].text = f"{datetime.now().strftime('%Y年%m月%d日')}現在"
    set_run(paras[3], f"氏名　{d.get('name','')}")
    set_run(paras[5], d.get("summary",""))
    set_run(paras[7], d.get("skills",""))

    companies = d.get("companies", [])

    # 1社目：テンプレートのテーブルに流し込む
    if companies:
        c = companies[0]
        set_run(paras[9],  f"（１）{c.get('company_name','')}")
        set_run(paras[10], f"・事業内容：{c.get('business','')}")
        set_run(paras[11], f"・資本金：{c.get('capital','')}万円")
        set_run(paras[12], f"・従業員数：{c.get('employees','')}名")

        t = doc.tables[0]
        cell_period = t.rows[1].cells[0]
        for para in cell_period.paragraphs:
            for run in para.runs: run.text = ""
        if cell_period.paragraphs:
            cell_period.paragraphs[0].add_run(
                f"{c.get('period_start','')} ～ {c.get('period_end','')}")

        cell_emp = t.rows[1].cells[1]
        for para in cell_emp.paragraphs:
            for run in para.runs: run.text = ""
        if cell_emp.paragraphs:
            cell_emp.paragraphs[0].add_run(
                f"雇用形態：{c.get('employment_type','正社員')}　配属先：{c.get('department','')}")

        cell_content = t.rows[2].cells[1]
        for para in cell_content.paragraphs:
            for run in para.runs: run.text = ""
        if cell_content.paragraphs:
            cell_content.paragraphs[0].add_run(
                f"《業務内容》\n{c.get('job_content','')}\n\n《実績》\n{c.get('achievement','')}")

    # テンプレート内の「自己PR」「以上」段落を削除
    paras_to_remove = []
    for p in doc.paragraphs:
        if p.text.strip() in ["自己PR", "以上"]:
            paras_to_remove.append(p)
    for p in paras_to_remove:
        p._element.getparent().remove(p._element)

    # 2社目以降：ドキュメント末尾にテーブル形式で追加する
    def add_company_block(doc, idx, c):
        # 会社名見出し
        p_head = doc.add_paragraph()
        run = p_head.add_run(f"（{idx}）{c.get('company_name','')}")
        run.bold = True
        run.font.size = Pt(10.5)

        # 会社情報
        doc.add_paragraph(f"・事業内容：{c.get('business','')}")
        doc.add_paragraph(f"・資本金：{c.get('capital','')}万円")
        doc.add_paragraph(f"・従業員数：{c.get('employees','')}名")

        # テーブル作成
        tbl = doc.add_table(rows=3, cols=2)
        try:
            tbl.style = "Table Grid"
        except Exception:
            pass

        # ヘッダー行
        tbl.rows[0].cells[0].text = "期間"
        tbl.rows[0].cells[1].text = "業務内容"

        # 期間・雇用形態
        tbl.rows[1].cells[0].text = f"{c.get('period_start','')} ～ {c.get('period_end','')}"
        tbl.rows[1].cells[1].text = f"雇用形態：{c.get('employment_type','正社員')}　配属先：{c.get('department','')}"

        # 業務内容・実績
        tbl.rows[2].cells[0].text = ""
        tbl.rows[2].cells[1].text = (
            f"《業務内容》\n{c.get('job_content','')}\n\n《実績》\n{c.get('achievement','')}"
        )

    for i, c in enumerate(companies[1:], start=2):
        if not c.get("company_name"):
            continue
        add_company_block(doc, i, c)

    # 自己PRを最後に追加
    p_pr_head = doc.add_paragraph()
    run = p_pr_head.add_run("自己PR")
    run.bold = True
    run.font.size = Pt(10.5)
    doc.add_paragraph(d.get("pr",""))
    doc.add_paragraph("以上")

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out

# ── プレビュー表示関数 ────────────────────────────────
def show_preview(d):
    """応募者データをプレビュー表示する"""
    st.markdown(f"**【氏名】** {d.get('furigana_name','')} / {d.get('name','')}　**【性別】** {d.get('gender','')}")
    st.markdown(f"**【生年月日】** {d.get('birthday','')}（{d.get('age','')}歳）")
    st.markdown(f"**【住所】** 〒{d.get('postal','')}　{d.get('address','')}　（{d.get('furigana_address','')}）")
    st.markdown(f"**【電話】** {d.get('phone','')}　**【メール】** {d.get('email','')}")
    st.markdown(f"**【最寄り駅】** {d.get('nearest_station','')}")
    st.markdown(f"**【扶養家族】** {d.get('dependents','')}人　**【配偶者】** {d.get('spouse','')}　**【配偶者扶養義務】** {d.get('spouse_support','')}")

    st.markdown("---")
    st.markdown("**【学歴】**")
    for e in d.get("education", []):
        if e.get("content"):
            st.markdown(f"　{e.get('year','')}年{e.get('month','')}月　{e.get('content','')}")

    st.markdown("**【職歴】**")
    for c in d.get("career", []):
        if c.get("content"):
            st.markdown(f"　{c.get('year','')}年{c.get('month','')}月　{c.get('content','')}")

    st.markdown("**【免許・資格】**")
    for l in d.get("licenses", []):
        if l.get("content"):
            st.markdown(f"　{l.get('year','')}年{l.get('month','')}月　{l.get('content','')}")

    st.markdown("---")
    st.markdown("**【自己PR】**")
    st.markdown(d.get("pr",""))

    st.markdown("---")
    st.markdown("**【職務要約】**")
    st.markdown(d.get("summary",""))

    st.markdown("**【活かせるスキル】**")
    st.markdown(d.get("skills",""))

    st.markdown("**【職務経歴】**")
    for i, c in enumerate(d.get("companies", [])):
        if not c.get("company_name"):
            continue
        st.markdown(f"**（{i+1}）{c.get('company_name','')}**　{c.get('period_start','')}〜{c.get('period_end','')}（{c.get('period_years','')}）")
        st.markdown(f"　事業内容：{c.get('business','')}　／　資本金：{c.get('capital','')}万円　／　従業員数：{c.get('employees','')}名")
        st.markdown(f"　雇用形態：{c.get('employment_type','')}　配属先：{c.get('department','')}")
        if c.get("job_content"):
            st.markdown(f"　業務内容：{c.get('job_content','')}")
        if c.get("achievement"):
            st.markdown(f"　実績：{c.get('achievement','')}")


# ── ヘルパー関数 ─────────────────────────────────────
def get_index(options, value):
    """valueがoptionsに存在すればそのindex、なければ0を返す"""
    try:
        return options.index(str(value)) if str(value) in options else 0
    except Exception:
        return 0

def parse_birthday(birthday_str):
    """'1995年4月10日' → ('1995','4','10')"""
    import re
    m = re.match(r"(\d+)年(\d+)月(\d+)日", str(birthday_str))
    if m:
        return m.group(1), m.group(2), m.group(3)
    return "", "", ""

def parse_yearmonth(ym_str):
    """'2020年4月' → ('2020','4')"""
    import re
    m = re.match(r"(\d+)年(\d+)月", str(ym_str))
    if m:
        return m.group(1), m.group(2)
    return "", ""

def parse_period(period_str):
    """'3年6ヶ月' → ('3','6')、'5年' → ('5','')"""
    import re
    y_m = re.match(r"(\d+)年(\d+)ヶ月", str(period_str))
    if y_m:
        return y_m.group(1), y_m.group(2)
    y_only = re.match(r"(\d+)年", str(period_str))
    if y_only:
        return y_only.group(1), ""
    m_only = re.match(r"(\d+)ヶ月", str(period_str))
    if m_only:
        return "", m_only.group(1)
    return "", ""


# ══════════════════════════════════════════════════════
# URLパラメータからclient_idを取得
# ══════════════════════════════════════════════════════
query_params = st.query_params
client_id = query_params.get("client", "admin")

# ══════════════════════════════════════════════════════
# サイドバー：モード選択
# ══════════════════════════════════════════════════════
mode = st.sidebar.radio("モード", ["📝 入力フォーム（ユーザー）", "🔐 管理画面"])

# ══════════════════════════════════════════════════════
# ユーザー側：入力フォーム
# ══════════════════════════════════════════════════════
if mode == "📝 入力フォーム（ユーザー）":
    st.title("📄 経歴入力フォーム")

    # 個人情報の注意書き
    st.info("🔒 入力された情報は、履歴書・職務経歴書作成の目的で使用されます。デモ利用時は架空の情報でお試しください。")
    st.markdown("---")

    # 確認画面から戻ったときのデータ復元
    _prev = st.session_state.get("preview_data", {})

    with st.form("resume_form"):

        st.subheader("① 氏名")
        furigana_name = st.text_input("ふりがな", value=_prev.get("furigana_name",""), placeholder="やまだ たろう")
        name = st.text_input("氏名 *", value=_prev.get("name",""), placeholder="山田 太郎")
        GENDER_OPT = ["", "男", "女"]
        gender = st.selectbox("性別", GENDER_OPT, index=get_index(GENDER_OPT, _prev.get("gender","")))

        st.subheader("② 生年月日")
        bd_col1, bd_col2, bd_col3 = st.columns(3)
        BD_YEAR_OPT = [""] + [str(y) for y in range(datetime.now().year - 15, 1919, -1)]
        BD_MON_OPT  = [""] + [str(m) for m in range(1, 13)]
        BD_DAY_OPT  = [""] + [str(d) for d in range(1, 32)]
        _pby, _pbm, _pbd = parse_birthday(_prev.get("birthday",""))
        bd_y = bd_col1.selectbox("年", BD_YEAR_OPT, index=get_index(BD_YEAR_OPT, _pby), key="bd_y")
        bd_m = bd_col2.selectbox("月", BD_MON_OPT,  index=get_index(BD_MON_OPT, _pbm),  key="bd_m")
        bd_d = bd_col3.selectbox("日", BD_DAY_OPT,  index=get_index(BD_DAY_OPT, _pbd),  key="bd_d")
        birthday = f"{bd_y}年{bd_m}月{bd_d}日" if bd_y and bd_m and bd_d else ""

        # 年齢自動計算
        if bd_y and bd_m and bd_d:
            try:
                today = datetime.now()
                born = datetime(int(bd_y), int(bd_m), int(bd_d))
                age_val = today.year - born.year - ((today.month, today.day) < (born.month, born.day))
                age = str(age_val)
                st.caption(f"年齢：{age}歳（自動計算）")
            except Exception:
                age = ""
        else:
            age = ""

        st.subheader("③ 住所")
        postal   = st.text_input("郵便番号", value=_prev.get("postal",""), placeholder="123-4567")
        address  = st.text_area("住所", value=_prev.get("address",""), placeholder="東京都新宿区〇〇1-2-3", height=68)
        furigana_address = st.text_input("住所ふりがな", value=_prev.get("furigana_address",""), placeholder="とうきょうとしんじゅくく〇〇")

        st.subheader("④ 電話番号")
        phone = st.text_input("電話番号", value=_prev.get("phone",""), placeholder="090-1234-5678")

        st.subheader("⑤ メールアドレス")
        email = st.text_input("メールアドレス", value=_prev.get("email",""), placeholder="example@email.com")

        # 年・月の選択肢
        YEAR_OPTIONS  = [""] + [str(y) for y in range(datetime.now().year, 1959, -1)]
        MONTH_OPTIONS = [""] + [str(m) for m in range(1, 13)]

        st.subheader("⑥ 学歴（高校から・西暦）")
        st.caption("年・月をプルダウンで選び、内容を入力してください")
        education = []
        _prev_edu = _prev.get("education", [{}]*6)
        cols_h = st.columns([1,1,4])
        cols_h[0].markdown("**年**")
        cols_h[1].markdown("**月**")
        cols_h[2].markdown("**内容**")
        for i in range(6):
            c1,c2,c3 = st.columns([1,1,4])
            _pe = _prev_edu[i] if i < len(_prev_edu) else {}
            y = c1.selectbox("年", YEAR_OPTIONS, index=get_index(YEAR_OPTIONS, _pe.get("year","")), key=f"edu_y{i}", label_visibility="collapsed")
            m = c2.selectbox("月", MONTH_OPTIONS, index=get_index(MONTH_OPTIONS, _pe.get("month","")), key=f"edu_m{i}", label_visibility="collapsed")
            c = c3.text_input("内容", value=_pe.get("content",""), key=f"edu_c{i}", placeholder="〇〇高等学校 卒業", label_visibility="collapsed")
            education.append({"year":y,"month":m,"content":c})

        st.subheader("⑦ 職歴")
        career = []
        _prev_car = _prev.get("career", [{}]*8)
        cols_h2 = st.columns([1,1,4])
        cols_h2[0].markdown("**年**")
        cols_h2[1].markdown("**月**")
        cols_h2[2].markdown("**内容**")
        for i in range(8):
            c1,c2,c3 = st.columns([1,1,4])
            _pc2 = _prev_car[i] if i < len(_prev_car) else {}
            y = c1.selectbox("年", YEAR_OPTIONS, index=get_index(YEAR_OPTIONS, _pc2.get("year","")), key=f"car_y{i}", label_visibility="collapsed")
            m = c2.selectbox("月", MONTH_OPTIONS, index=get_index(MONTH_OPTIONS, _pc2.get("month","")), key=f"car_m{i}", label_visibility="collapsed")
            c = c3.text_input("内容", value=_pc2.get("content",""), key=f"car_c{i}", placeholder="〇〇株式会社 入社", label_visibility="collapsed")
            career.append({"year":y,"month":m,"content":c})

        st.subheader("⑧ 免許・資格")
        licenses = []
        _prev_lic = _prev.get("licenses", [{}]*5)
        cols_h3 = st.columns([1,1,4])
        cols_h3[0].markdown("**年**")
        cols_h3[1].markdown("**月**")
        cols_h3[2].markdown("**内容**")
        for i in range(5):
            c1,c2,c3 = st.columns([1,1,4])
            _pl = _prev_lic[i] if i < len(_prev_lic) else {}
            y = c1.selectbox("年", YEAR_OPTIONS, index=get_index(YEAR_OPTIONS, _pl.get("year","")), key=f"lic_y{i}", label_visibility="collapsed")
            m = c2.selectbox("月", MONTH_OPTIONS, index=get_index(MONTH_OPTIONS, _pl.get("month","")), key=f"lic_m{i}", label_visibility="collapsed")
            c = c3.text_input("内容", value=_pl.get("content",""), key=f"lic_c{i}", placeholder="普通自動車第一種運転免許 取得", label_visibility="collapsed")
            licenses.append({"year":y,"month":m,"content":c})

        st.subheader("⑨ 自己PR")
        st.caption("どんな人物か・何を意識してきたか・どんなふうに頑張れるかを書いてください")
        pr = st.text_area("自己PR *", value=_prev.get("pr",""), height=200,
                          placeholder="コミュニケーションを大切にし、チームで目標達成することを意識してきました。")

        st.subheader("⑩ 最寄り駅")
        nearest_station = st.text_input("最寄り駅", value=_prev.get("nearest_station",""), placeholder="〇〇線 〇〇駅")

        st.subheader("⑪ 家族情報")
        fc1,fc2,fc3 = st.columns(3)
        dependents     = fc1.text_input("扶養家族（人数）", value=_prev.get("dependents",""), placeholder="0")
        SPOUSE_OPT = ["無","有"]
        spouse         = fc2.selectbox("配偶者", SPOUSE_OPT, index=get_index(SPOUSE_OPT, _prev.get("spouse","無")))
        spouse_support = fc3.selectbox("配偶者の扶養義務", SPOUSE_OPT, index=get_index(SPOUSE_OPT, _prev.get("spouse_support","無")))

        st.markdown("---")
        st.header("📋 職務経歴書")
        summary = st.text_area("職務要約 *", value=_prev.get("summary",""), height=120,
                               placeholder="〇〇業界で△年間の経験を持ちます。")
        skills  = st.text_area("活かせるスキル", value=_prev.get("skills",""), height=100,
                               placeholder="・〇〇スキル\n・△△の経験")

        st.subheader("職務経歴（最大5社）")
        companies = []
        _prev_comp = _prev.get("companies", [{}]*MAX_COMPANIES)
        for i in range(MAX_COMPANIES):
            _pc = _prev_comp[i] if i < len(_prev_comp) else {}
            with st.expander(f"（{i+1}）会社情報", expanded=(i==0)):
                cname = st.text_input("会社名", value=_pc.get("company_name",""), key=f"cn{i}", placeholder="〇〇株式会社")

                # 入社年月（プルダウン）
                st.caption("入社年月")
                ps_col1, ps_col2 = st.columns(2)
                _psy, _psm = parse_yearmonth(_pc.get("period_start",""))
                ps_y = ps_col1.selectbox("入社年", YEAR_OPTIONS, index=get_index(YEAR_OPTIONS, _psy), key=f"ps_y{i}", label_visibility="collapsed")
                ps_m = ps_col2.selectbox("入社月", MONTH_OPTIONS, index=get_index(MONTH_OPTIONS, _psm), key=f"ps_m{i}", label_visibility="collapsed")
                ps = f"{ps_y}年{ps_m}月" if ps_y and ps_m else ""

                # 退社年月（プルダウン）
                st.caption("退社年月（在籍中の場合は空欄）")
                pe_col1, pe_col2 = st.columns(2)
                _pend = _pc.get("period_end","")
                _pey, _pem = parse_yearmonth(_pend) if _pend != "現在" else ("","")
                pe_y = pe_col1.selectbox("退社年", YEAR_OPTIONS, index=get_index(YEAR_OPTIONS, _pey), key=f"pe_y{i}", label_visibility="collapsed")
                pe_m = pe_col2.selectbox("退社月", MONTH_OPTIONS, index=get_index(MONTH_OPTIONS, _pem), key=f"pe_m{i}", label_visibility="collapsed")
                pe = f"{pe_y}年{pe_m}月" if pe_y and pe_m else "現在"

                # 在籍期間（年・月）
                st.caption("在籍期間")
                py_col1, py_col2 = st.columns(2)
                TENURE_YEAR_OPT  = [""] + [str(n) for n in range(0, 51)]
                TENURE_MON_OPT   = [""] + [str(n) for n in range(0, 12)]
                _tyy, _tym = parse_period(_pc.get("period_years",""))
                py_y = py_col1.selectbox("○○年", TENURE_YEAR_OPT, index=get_index(TENURE_YEAR_OPT, _tyy), key=f"py_y{i}")
                py_m = py_col2.selectbox("○○ヶ月", TENURE_MON_OPT, index=get_index(TENURE_MON_OPT, _tym), key=f"py_m{i}")
                py_parts = []
                if py_y: py_parts.append(f"{py_y}年")
                if py_m: py_parts.append(f"{py_m}ヶ月")
                py = "".join(py_parts)

                biz  = st.text_input("事業内容", value=_pc.get("business",""), key=f"biz{i}", placeholder="〇〇の製造・販売")
                cc4,cc5 = st.columns(2)
                cap  = cc4.text_input("資本金（万円）", value=_pc.get("capital",""), key=f"cap{i}", placeholder="5000")
                emp  = cc5.text_input("従業員数（名）", value=_pc.get("employees",""), key=f"emp{i}", placeholder="200")
                ETYPE_OPT = ["正社員","契約社員","アルバイト","派遣","その他"]
                etype = st.selectbox("雇用形態", ETYPE_OPT, index=get_index(ETYPE_OPT, _pc.get("employment_type","正社員")), key=f"et{i}")
                dept  = st.text_input("配属先", value=_pc.get("department",""), key=f"dept{i}", placeholder="営業部 第一課")
                jc    = st.text_area("業務内容", value=_pc.get("job_content",""), key=f"jc{i}", height=100,
                                     placeholder="・〇〇業務を担当")
                ach   = st.text_area("実績", value=_pc.get("achievement",""), key=f"ach{i}", height=80,
                                     placeholder="・〇〇を達成し前年比△%向上に貢献")
                companies.append({
                    "company_name":cname,"period_start":ps,"period_end":pe,
                    "period_years":py,"business":biz,"capital":cap,
                    "employees":emp,"employment_type":etype,
                    "department":dept,"job_content":jc,"achievement":ach,
                })

        st.markdown("---")
        submitted = st.form_submit_button("📋 内容を確認する", type="primary", use_container_width=True)

    if submitted:
        errors = []
        if not name.strip():    errors.append("氏名を入力してください。")
        if not pr.strip():      errors.append("自己PRを入力してください。")
        if not summary.strip(): errors.append("職務要約を入力してください。")
        if email.strip() and ("@" not in email or "." not in email):
            errors.append("メールアドレスの形式が正しくありません（@ と . を含めてください）。")
        if phone.strip() and not all(c.isdigit() or c == "-" for c in phone.strip()):
            errors.append("電話番号は数字とハイフン（-）のみ入力してください。")
        if errors:
            for e in errors:
                st.error(f"❌ {e}")
        else:
            # session_stateに一時保存して確認画面へ
            st.session_state["preview_data"] = {
                "name":name,"furigana_name":furigana_name,
                "gender":gender,
                "client_id":client_id,
                "birthday":birthday,"age":age,
                "postal":postal,"address":address,"furigana_address":furigana_address,
                "phone":phone,"email":email,
                "education":education,"career":career,"licenses":licenses,
                "pr":pr,"nearest_station":nearest_station,
                "dependents":dependents,"spouse":spouse,"spouse_support":spouse_support,
                "summary":summary,"skills":skills,"companies":companies,
            }
            st.session_state["show_preview"] = True
            st.rerun()

    # 確認画面
    if st.session_state.get("show_preview") and "preview_data" in st.session_state:
        d = st.session_state["preview_data"]
        st.markdown("---")
        st.subheader("📋 入力内容の確認")
        st.caption("内容をご確認ください。修正がある場合は「編集に戻る」を押してください。")

        show_preview(d)

        col_back, col_submit = st.columns(2)
        if col_back.button("✏️ 編集に戻る", use_container_width=True):
            st.session_state["show_preview"] = False
            st.rerun()

        if col_submit.button("✅ このまま送信する", type="primary", use_container_width=True):
            record = {
                "id": datetime.now().strftime("%Y%m%d%H%M%S") + "_" + uuid.uuid4().hex[:8],
                "submitted_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                **d,
            }
            records = load_data()
            records.append(record)
            save_data(records)
            save_to_gsheet_full(record)
            st.session_state["show_preview"] = False
            st.session_state.pop("preview_data", None)
            st.success("✅ 送信が完了しました！担当アドバイザーが確認します。")
            st.balloons()

# ══════════════════════════════════════════════════════
# 管理者側：管理画面
# ══════════════════════════════════════════════════════
else:
    st.title("🔐 管理画面")

    if "admin_logged_in" not in st.session_state:
        st.session_state.admin_logged_in = False

    if not st.session_state.admin_logged_in:
        pw = st.text_input("パスワード", type="password")
        if st.button("ログイン"):
            # client_idに対応するパスワードを確認
            correct_pass = CLIENT_PASSWORDS.get(client_id, CLIENT_PASSWORDS.get("admin", "admin1234"))
            if pw == correct_pass:
                st.session_state.admin_logged_in = True
                st.session_state.admin_client_id = client_id
                st.rerun()
            else:
                st.error("❌ パスワードが違います")
    else:
        st.caption("⚠️ 本番利用時は管理パスワードを変更してください（st.secrets を使用）")
        st.success("✅ ログイン中")
        if st.button("ログアウト"):
            st.session_state.admin_logged_in = False
            st.rerun()

        logged_client = st.session_state.get("admin_client_id", "admin")
        records = load_all_from_gsheet(client_id=logged_client)
        if not records:
            # フォールバック：ローカルファイルから読み込み
            records = load_data()
        if not records:
            st.info("まだ送信されたデータがありません。")
        else:
            st.subheader(f"📋 送信一覧（{len(records)}件）")

            # 応募者一覧CSVダウンロード
            csv_cols = ["id","submitted_at","name","birthday","age",
                        "phone","email","nearest_station","summary"]
            df_list = []
            for r in records:
                df_list.append({c: r.get(c,"") for c in csv_cols})
            df_csv = pd.DataFrame(df_list, columns=csv_cols)
            csv_bytes = df_csv.to_csv(index=False, encoding="utf-8-sig").encode("utf-8-sig")
            st.download_button(
                "⬇️ 応募者一覧CSVをダウンロード",
                data=csv_bytes,
                file_name=f"応募者一覧_{datetime.now().strftime('%Y%m%d')}.csv",
                mime="text/csv",
            )

            st.markdown("---")

            for r in reversed(records):
                with st.expander(f"📄 {r.get('name','')}　送信日時：{r.get('submitted_at','')}"):
                    col1, col2 = st.columns(2)

                    # 履歴書ダウンロード
                    try:
                        buf_r = make_rirekisho(r)
                        col1.download_button(
                            "⬇️ 履歴書.docxをダウンロード",
                            data=buf_r,
                            file_name=f"履歴書_{r.get('name','')}_{r.get('id','')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"r_{r.get('id','')}",
                            use_container_width=True,
                        )
                    except FileNotFoundError as e:
                        col1.error(str(e))

                    # 職務経歴書ダウンロード
                    try:
                        buf_s = make_shokumu(r)
                        col2.download_button(
                            "⬇️ 職務経歴書.docxをダウンロード",
                            data=buf_s,
                            file_name=f"職務経歴書_{r.get('name','')}_{r.get('id','')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"s_{r.get('id','')}",
                            use_container_width=True,
                        )
                    except FileNotFoundError as e:
                        col2.error(str(e))

                    st.markdown("---")
                    st.subheader("📋 入力内容プレビュー")
                    show_preview(r)

                    # 削除機能
                    st.markdown("---")
                    del_check = st.checkbox(
                        "このデータを削除する（チェックを入れてから削除ボタンを押してください）",
                        key=f"del_check_{r.get('id','')}"
                    )
                    if del_check:
                        if st.button("🗑️ 削除する", key=f"del_{r.get('id','')}", type="secondary"):
                            # スプレッドシートから削除
                            delete_from_gsheet(r.get("id",""))
                            # ローカルファイルからも削除
                            records_new = load_data()
                            records_new = [x for x in records_new if x.get("id") != r.get("id")]
                            save_data(records_new)
                            st.success(f"✅ 「{r.get('name','')}」のデータを削除しました。")
                            st.rerun()
